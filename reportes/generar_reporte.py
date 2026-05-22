"""
generar_reporte.py
==================
Inserta todas las gráficas de output/ en el Word 17-CRT_DGP_017_2026.docx,
preservando el header y footer originales del documento.

Estructura del documento generado:
  1. Portada
  2. Índice de figuras
  3. Una página por figura (título + imagen)
  4. Conclusiones

Uso:
    py generar_reporte.py
"""

import os
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Rutas ────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(r"C:\Users\ivan-\Documents\GitHub\anuario")
OUTPUT_DIR = BASE_DIR / "output"
SRC_DOCX   = BASE_DIR / "17-CRT_DGP_017_2026.docx"
OUT_DOCX   = BASE_DIR / "17-CRT_DGP_017_2026_reporte.docx"

# ── Títulos de cada figura ────────────────────────────────────────────────────
# Mapa: nombre_de_archivo (sin ext, case-insensitive key) -> título descriptivo
TITULOS = {
    "figura_a1":   "Figura A.1 — Producto Interno Bruto (PIB) y contribución del PIB de los subsectores de Telecomunicaciones y Radiodifusión",
    "figura_a2":   "Figura A.2 — Personal ocupado en los sectores de Telecomunicaciones y Radiodifusión",
    "figura_a3":   "Figura A.3 — Índice Nacional de Precios al Consumidor (INPC) y de las Comunicaciones",
    "figura_a4":   "Figura A.4 — Inversión en los sectores de Telecomunicaciones y Radiodifusión",
    "figura_a5":   "Figura A.5 — Inversión Extranjera Directa (IED) en Telecomunicaciones",
    "figura_a6":   "Figura A.6 — Ingresos y egresos del sector de Telecomunicaciones",
    "figura_a7":   "Figura A.7 — Hogares con acceso a servicios de Telecomunicaciones fijas por decil de ingreso",
    "figura_a8":   "Figura A.8 — Gasto en servicios de Telecomunicaciones fijas por decil de ingreso",
    "figura_a9":   "Figura A.9 — Hogares con servicios de Telecomunicaciones móviles por decil de ingreso",
    "figura_a10":  "Figura A.10 — Gasto en servicios de Telecomunicaciones móviles por decil de ingreso",
    "figura_b1":   "Figura B.1 — Hogares con servicios fijos de Telecomunicaciones (Internet, TV restringida y Telefonía fija)",
    "figura_b4":   "Figura B.4 — Líneas del servicio fijo de Telefonía (2000–2023)",
    "figura_b5":   "Figura B.5 — Penetración del servicio fijo de Telefonía por cada 100 hogares",
    "figura_b6":   "Figura B.6 — Penetración residencial del servicio fijo de Telefonía por entidad federativa",
    "figura_b7":   "Figura B.7 — Penetración no residencial del servicio fijo de Telefonía por entidad federativa",
    "figura_b8":   "Figura B.8 — Tráfico del servicio fijo de Telefonía (millones de minutos)",
    "figura_b9":   "Figura B.9 — Participación de mercado del servicio fijo de Telefonía",
    "figura_b10":  "Figura B.10 — Índice Herfindahl-Hirschman (IHH) del servicio fijo de Internet (BAF)",
    "figura_b11":  "Figura B.11 — Accesos del servicio fijo de Internet (2000–2023)",
    "figura_b12":  "Figura B.12 — Penetración del servicio fijo de Internet por cada 100 hogares",
    "figura_b13":  "Figura B.13 — Penetración del servicio fijo de Internet por entidad federativa (residencial)",
    "figura_b14":  "Figura B.14 — Penetración del servicio fijo de Internet por entidad federativa (no residencial)",
    "figura_b15":  "Figura B.15 — Velocidad promedio de descarga de Internet fijo por entidad federativa",
    "figura_b16":  "Figura B.16 — Suscriptores del servicio de TV restringida",
    "figura_b17":  "Figura B.17 — Participación de mercado del servicio de TV restringida",
    "figura_b18":  "Figura B.18 — IHH del servicio de TV restringida",
    "figura_b19":  "Figura B.19 — Penetración del servicio de TV restringida por cada 100 hogares",
    "figura_b20":  "Figura B.20 — Penetración del servicio de TV restringida por entidad federativa",
    "figura_b21":  "Figura B.21 — Líneas del servicio móvil (2000–2023)",
    "figura_b22":  "Figura B.22 — Penetración del servicio móvil por cada 100 habitantes",
    "figura_b23":  "Figura B.23 — Participación de mercado del servicio móvil",
    "figura_b24":  "Figura B.24 — IHH del servicio móvil",
    "figura_b25":  "Figura B.25 — Tráfico de voz del servicio móvil",
    "figura_b2":   "Figura B.2 — Mapa de cobertura de red móvil (2G/3G/4G)",
    "figura_b3":   "Figura B.3 — Mapa de cobertura de red móvil LTE/4G",
    "figura_c1":   "Figura C.1 — Concesionarios y autorizados del servicio fijo de Telefonía",
    "figura_c2":   "Figura C.2 — Concesionarios y autorizados del servicio fijo de Internet",
    "figura_c3":   "Figura C.3 — Concesionarios y autorizados del servicio móvil",
    "figura_c4":   "Figura C.4 — Cobertura de la red de banda ancha fija por municipio",
    "figura_c5":   "Figura C.5 — Cobertura de la red móvil por municipio",
    "figura_c6":   "Figura C.6 — Infraestructura de torres de telecomunicaciones por entidad",
    "figura_c7":   "Figura C.7 — Distribución de espectro radioeléctrico concesionado",
    "figura_c8":   "Figura C.8 — Espectro radioeléctrico concesionado por banda de frecuencia",
    "figura_c9":   "Figura C.9 — Espectro asignado por servicio y operador",
    "figura_c10":  "Figura C.10 — Concesiones de espectro vigentes",
    "figura_c11":  "Figura C.11 — Interconexión entre operadores de telecomunicaciones",
    "figura_c12":  "Figura C.12 — Tráfico de interconexión por tipo de servicio",
    "figura_c13":  "Figura C.13 — Números portados acumulados (NNP)",
    "figura_c14":  "Figura C.14 — Portabilidad numérica: solicitudes mensuales",
    "figura_c15":  "Figura C.15 — Portabilidad numérica por entidad federativa",
    "figura_c16":  "Figura C.16 — Resolución de disputas en telecomunicaciones",
    "figura_d1":   "Figura D.1 — Estaciones de radio AM y FM concesionadas y permisionadas",
    "figura_d2":   "Figura D.2 — Estaciones de televisión abierta concesionadas y permisionadas",
    "figura_d3":   "Figura D.3 — Distribución de estaciones de radio por entidad federativa",
    "figura_d4":   "Figura D.4 — Distribución de estaciones de TV abierta por entidad federativa",
    "figura_d5":   "Figura D.5 — Audiencia de radio por entidad federativa",
    "figura_d6":   "Figura D.6 — Audiencia de televisión abierta por entidad federativa",
    "figura_d7":   "Figura D.7 — Contenidos producidos en México: horas transmitidas",
    "figura_d8":   "Figura D.8 — Contenidos de producción nacional vs extranjera en TV",
    "figura_d9":   "Figura D.9 — Contenidos de TV abierta por género programático",
    "figura_d10":  "Figura D.10 — Horario de transmisión en TV abierta",
    "figura_d11":  "Figura D.11 — Inversión publicitaria en medios de comunicación",
    "figura_e1":   "Figura E.1 — Suscriptores de servicios OTT de video",
    "figura_e2":   "Figura E.2 — Comparativo de suscriptores OTT vs TV restringida",
    "figura_e3":   "Figura E.3 — Uso de plataformas digitales por tipo de servicio",
    "figura_e4":   "Figura E.4 — Comercio electrónico: valor de transacciones",
    "figura_e5":   "Figura E.5 — Adopción de servicios digitales por segmento de usuario",
    "figura_e6":   "Figura E.6 — Seguridad de la información: incidentes reportados",
    "figura_e7":   "Figura E.7 — Mapa de infraestructura de fibra óptica nacional",
    "figura_e8":   "Figura E.8 — Velocidades de Internet fijo: distribución por tecnología",
    "figura_e9":   "Figura E.9 — Velocidades de Internet móvil: distribución por tecnología",
    "figura_f1_1": "Figura F.1.1 — Usuarios de Internet por entidad federativa (total)",
    "figura_f1_2": "Figura F.1.2 — Usuarios de Internet: distribución por sexo y edad",
    "figura_f1_3": "Figura F.1.3 — Usuarios de Internet: distribución por nivel educativo",
    "figura_f1_4": "Figura F.1.4 — Usuarios de Internet: distribución por actividad principal",
    "figura_f2":   "Figura F.2 — Dispositivos utilizados para acceder a Internet",
    "figura_f4":   "Figura F.4 — Motivos de uso de Internet por tipo de actividad",
    "figura_f5":   "Figura F.5 — Usuarios de redes sociales y mensajería instantánea",
    "figura_f6":   "Figura F.6 — Frecuencia de uso de Internet por grupo de edad",
    "figura_f7":   "Figura F.7 — Lugar de acceso a Internet",
    "figura_f8":   "Figura F.8 — Percepción de seguridad en el uso de Internet",
    "figura_f9":   "Figura F.9 — Barreras para el acceso a Internet",
    "figura_f10":  "Figura F.10 — Uso de gobierno electrónico (e-gobierno)",
    "figura_f11":  "Figura F.11 — Usuarios de comercio electrónico",
    "figura_f12":  "Figura F.12 — Usuarios de banca en línea",
    "figura_f13":  "Figura F.13 — Uso de servicios de salud digital",
    "figura_f14":  "Figura F.14 — Usuarios de educación en línea",
    "figura_f15":  "Figura F.15 — Habilidades digitales de la población",
    "figura_f16":  "Figura F.16 — Brecha digital: comparativo rural vs urbano",
    "figura_g1":   "Figura G.1 — Indicadores internacionales de telecomunicaciones: México vs OCDE",
    "figura_h1":   "Figura H.1 — (Datos parciales) Infraestructura de red de nueva generación",
    "figura_h2":   "Figura H.2 — (Datos parciales) Despliegue de redes 5G en México",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def _key(filename: str) -> str:
    """Normaliza el nombre de archivo a clave del diccionario TITULOS."""
    name = Path(filename).stem.lower()
    # normalizar separadores: Figura_A1 -> figura_a1, figura_f1.1 -> figura_f1_1
    name = name.replace(".", "_")
    return name


def get_titulo(png_path: Path) -> str:
    """Devuelve el título descriptivo de una figura, o el nombre del archivo si no se conoce."""
    key = _key(png_path.name)
    if key in TITULOS:
        return TITULOS[key]
    # Intentar con variante sin guión bajo (ej. figura_b10 -> figurab10)
    key2 = key.replace("_", "")
    for k, v in TITULOS.items():
        if k.replace("_", "") == key2:
            return v
    # Fallback: derivar título del nombre de archivo
    return png_path.stem.replace("_", " ").title()


def orden_figura(png_path: Path) -> tuple:
    """
    Clave de ordenamiento: sección (A-H) numérica, luego número principal,
    luego subnúmero (para F.1.1, F.1.2, etc.).
    """
    name = png_path.stem.lower()
    # extrae la sección (primera letra después de 'figura_')
    m = re.match(r"figura_([a-h])(\d+)(?:[._](\d+))?", name)
    if m:
        seccion = ord(m.group(1)) - ord('a')
        num1    = int(m.group(2))
        num2    = int(m.group(3)) if m.group(3) else 0
        return (seccion, num1, num2)
    return (99, 0, 0)


def recopilar_figuras() -> list[Path]:
    """Devuelve lista de PNGs del directorio output/, ordenada por figura."""
    pngs = [p for p in OUTPUT_DIR.glob("*.png")]
    pngs_filtradas = [p for p in pngs if re.search(r"figura", p.stem, re.IGNORECASE)]
    return sorted(pngs_filtradas, key=orden_figura)


# ── Copia de header/footer ────────────────────────────────────────────────────

def _clone_element(el):
    return copy.deepcopy(el)


def copy_header_footer(src_section, dst_section):
    """
    Copia el header y footer de src_section a dst_section a nivel XML.
    """
    for part_name in ("header", "footer"):
        src_part = getattr(src_section, f"{part_name}Reference", None)

        # Acceso directo por XML
        for which in ("default", "first", "even"):
            try:
                src_hdr = getattr(src_section, f"{which}_{part_name}")
                dst_hdr = getattr(dst_section, f"{which}_{part_name}")
                if src_hdr and src_hdr._element is not None:
                    # Reemplazar el contenido XML del destino con el del origen
                    src_xml = _clone_element(src_hdr._element)
                    dst_hdr._element.getparent().replace(dst_hdr._element, src_xml)
            except Exception:
                pass


def set_header_footer_linked(section, link: bool = True):
    """Vincula el header/footer de una sección al anterior (para que se hereden)."""
    sectPr = section._sectPr
    for tag in ("headerReference", "footerReference"):
        for ref in sectPr.findall(qn("w:" + tag)):
            sectPr.remove(ref)


# ── Formato de párrafos ───────────────────────────────────────────────────────

def set_run_font(run, size_pt=11, bold=False, color=None, font_name="Calibri"):
    run.font.name    = font_name
    run.font.size    = Pt(size_pt)
    run.font.bold    = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_centered_para(doc, text, size_pt=11, bold=False, color=None,
                       space_before=0, space_after=0, font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold, color=color, font_name=font_name)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br   # devolvemos el OxmlElement para usarlo directamente


def insert_page_break(doc):
    """Inserta un salto de página limpio."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ── Secciones del documento ───────────────────────────────────────────────────

CRT_VERDE   = (0, 104, 74)     # verde oscuro CRT (RGB)
CRT_GRIS    = (80, 80, 80)     # gris texto
NEGRO       = (0, 0, 0)

def build_portada(doc):
    """Agrega la sección de portada."""
    # Espaciado superior
    for _ in range(4):
        doc.add_paragraph()

    add_centered_para(doc,
        "Anuario Estadístico 2024 — IFT",
        size_pt=22, bold=True, color=CRT_VERDE, font_name="Calibri",
        space_after=6)

    add_centered_para(doc,
        "Reproducción de Gráficas con Python",
        size_pt=16, bold=False, color=CRT_GRIS, font_name="Calibri",
        space_after=4)

    # Línea separadora (usar guiones como sustituto visual)
    add_centered_para(doc, "─" * 48, size_pt=10, color=CRT_VERDE, space_after=12)

    add_centered_para(doc, "Alumno:", size_pt=12, bold=True, color=NEGRO)
    add_centered_para(doc, "[Nombre del Alumno]", size_pt=12, color=CRT_GRIS, space_after=8)

    add_centered_para(doc, "Asesor:", size_pt=12, bold=True, color=NEGRO)
    add_centered_para(doc, "Iván", size_pt=12, color=CRT_GRIS, space_after=8)

    add_centered_para(doc, "Fecha:", size_pt=12, bold=True, color=NEGRO)
    fecha_str = date.today().strftime("%d de %B de %Y")
    # Traducir mes al español
    MESES = {
        "January":"enero","February":"febrero","March":"marzo","April":"abril",
        "May":"mayo","June":"junio","July":"julio","August":"agosto",
        "September":"septiembre","October":"octubre","November":"noviembre","December":"diciembre",
    }
    for en, es in MESES.items():
        fecha_str = fecha_str.replace(en, es)
    add_centered_para(doc, fecha_str, size_pt=12, color=CRT_GRIS, space_after=8)

    add_centered_para(doc,
        "Proyecto: Reproducción de Gráficas del Anuario Estadístico 2024 del IFT",
        size_pt=10, color=CRT_GRIS, space_before=20)


def build_indice(doc, figuras: list[Path]):
    """Agrega el índice de figuras."""
    insert_page_break(doc)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("Índice de Figuras")
    set_run_font(run, size_pt=16, bold=True, color=CRT_VERDE)
    p_titulo.paragraph_format.space_after = Pt(12)

    for i, png in enumerate(figuras, start=1):
        titulo = get_titulo(png)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. {titulo}")
        set_run_font(run, size_pt=10, color=NEGRO)


def build_figura_pages(doc, figuras: list[Path]):
    """Agrega una página por figura: título + imagen."""
    for png in figuras:
        insert_page_break(doc)
        titulo = get_titulo(png)

        # Título de la figura
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(titulo)
        set_run_font(run, size_pt=12, bold=True, color=CRT_VERDE, font_name="Calibri")

        # Insertar imagen centrada
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after  = Pt(4)

        run_img = p_img.add_run()
        # Ajustar ancho máximo a 15 cm para que quepa en la página
        try:
            run_img.add_picture(str(png), width=Cm(15))
        except Exception as e:
            p_err = doc.add_paragraph()
            run_err = p_err.add_run(f"[Error cargando imagen: {png.name} — {e}]")
            set_run_font(run_err, size_pt=9, color=(200, 0, 0))


def build_conclusiones(doc):
    """Agrega la sección de conclusiones."""
    insert_page_break(doc)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_titulo.add_run("Conclusiones")
    set_run_font(run, size_pt=14, bold=True, color=CRT_VERDE)
    p_titulo.paragraph_format.space_after = Pt(10)

    parrafos = [
        (
            "Hasta ahora se han replicado las figuras de las secciones A (Indicadores "
            "económicos), B (Servicios fijos y móviles), C (Infraestructura y regulación), "
            "D (Radiodifusión), E (Economía digital), F (Uso de Internet y habilidades "
            "digitales) y G (Comparativos internacionales)."
        ),
        (
            "No se cuenta con los datos necesarios para generar las gráficas de la sección H, "
            "ya que corresponden a información de carácter privado que no está disponible en "
            "las fuentes de datos públicas consultadas (IFT BIT, INEGI, Secretaría de "
            "Economía, entre otras)."
        ),
        (
            "Se identificaron discrepancias entre los datos originales del Anuario y los "
            "calculados a partir de las fuentes públicas. Los datos descargados son datos "
            "crudos (sin procesar), mientras que el PDF del Anuario presenta datos procesados "
            "y redondeados por el IFT. El propio Anuario señala: \"La información reportada "
            "está sujeta a revisiones y a modificaciones derivadas de cambios que realizan "
            "los operadores a las cifras previamente reportadas\", lo que explica diferencias "
            "adicionales entre las versiones de datos."
        ),
        (
            "Las principales fuentes de discrepancia detectadas son: (1) revisiones "
            "posteriores a la publicación del Anuario por parte de los operadores de "
            "telecomunicaciones; (2) uso de versiones diferentes de proyecciones "
            "demográficas (CONAPO/INEGI); y (3) procesos de redondeo y agregación aplicados "
            "internamente por el IFT antes de publicar el Anuario."
        ),
    ]

    for texto in parrafos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(texto)
        set_run_font(run, size_pt=11, color=NEGRO)


# ── Programa principal ────────────────────────────────────────────────────────

def main():
    print(f"Abriendo plantilla: {SRC_DOCX}")
    doc = Document(str(SRC_DOCX))

    # Limpiar SOLO el contenido del body, preservando la estructura de sección
    # (headers/footers quedan intactos porque son partes relacionadas del docx)
    body = doc.element.body

    # Guardar el último sectPr (contiene márgenes, header/footer, etc.)
    last_sectPr = body.find(qn("w:sectPr"))
    if last_sectPr is None:
        # buscar en el último párrafo
        last_para = body.findall(qn("w:p"))[-1]
        pPr = last_para.find(qn("w:pPr"))
        if pPr is not None:
            last_sectPr = pPr.find(qn("w:sectPr"))

    # Eliminar todos los párrafos y tablas del body (no el sectPr raíz)
    children_to_remove = []
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "sectPr":
            children_to_remove.append(child)

    for child in children_to_remove:
        body.remove(child)

    # Si el sectPr estaba incrustado en un párrafo también se eliminó;
    # re-añadirlo al final del body para preservar headers/footers
    if last_sectPr is not None and last_sectPr.getparent() is None:
        body.append(last_sectPr)

    # Ahora agregar contenido
    figuras = recopilar_figuras()
    print(f"Se encontraron {len(figuras)} figuras en {OUTPUT_DIR}")

    print("  -> Generando portada...")
    build_portada(doc)

    print("  -> Generando indice de figuras...")
    build_indice(doc, figuras)

    print("  -> Insertando figuras...")
    build_figura_pages(doc, figuras)

    print("  -> Generando conclusiones...")
    build_conclusiones(doc)

    print(f"Guardando documento en: {OUT_DOCX}")
    doc.save(str(OUT_DOCX))
    print("OK! Listo!")


if __name__ == "__main__":
    main()
