"""
Orquestador de scripts del Anuario Estadístico 2024 del IFT.

Ejecuta todos los scripts de Python en las carpetas A–H,
genera un reporte en Word (.docx) y muestra un resumen
en terminal con archivos completados, con error y sin gráfica.

Uso:
    python orquestador.py
"""

import subprocess
import sys
import os
import time
import glob
import re
from pathlib import Path
from datetime import datetime

# ---------------------------------------------------------------------------
# Configuración
# ---------------------------------------------------------------------------
SCRIPTS_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPTS_DIR.parent
OUTPUT_DIR = PROJECT_ROOT / "output"
REPORT_DIR = PROJECT_ROOT / "reportes"

CARPETAS = ["A", "B", "C", "D", "E", "F", "G", "H"]

# Tamaño máximo (bytes) para considerar un script como "placeholder vacío"
# Los placeholders tienen ~180-185 bytes y sólo importan el logger.
PLACEHOLDER_MAX_SIZE = 250

# ---------------------------------------------------------------------------
# Utilidades de terminal con color (ANSI)
# ---------------------------------------------------------------------------
class Color:
    VERDE   = "\033[92m"
    ROJO    = "\033[91m"
    AMARILLO = "\033[93m"
    CYAN    = "\033[96m"
    BOLD    = "\033[1m"
    RESET   = "\033[0m"
    GRIS    = "\033[90m"


def titulo(texto: str) -> None:
    ancho = 70
    print(f"\n{Color.CYAN}{Color.BOLD}{'═' * ancho}")
    print(f"  {texto}")
    print(f"{'═' * ancho}{Color.RESET}\n")


def seccion(texto: str) -> None:
    print(f"\n{Color.BOLD}{Color.CYAN}── {texto} ──{Color.RESET}")


# ---------------------------------------------------------------------------
# Recopilar scripts
# ---------------------------------------------------------------------------
def recopilar_scripts() -> list[Path]:
    """Devuelve una lista ordenada de todos los archivos .py en las carpetas."""
    scripts = []
    for carpeta in CARPETAS:
        ruta = SCRIPTS_DIR / carpeta
        if not ruta.is_dir():
            print(f"{Color.AMARILLO}⚠  Carpeta no encontrada: {ruta}{Color.RESET}")
            continue
        archivos = sorted(ruta.glob("figura_*.py"))
        scripts.extend(archivos)
    return scripts


def es_placeholder(script: Path) -> bool:
    """Determina si un script es un placeholder vacío (sin lógica real)."""
    return script.stat().st_size <= PLACEHOLDER_MAX_SIZE


# ---------------------------------------------------------------------------
# Detectar gráficas generadas
# ---------------------------------------------------------------------------
def archivos_output_antes() -> set[str]:
    """Snapshot de los archivos en la carpeta output."""
    if not OUTPUT_DIR.exists():
        return set()
    return set(f.name for f in OUTPUT_DIR.iterdir() if f.is_file())


def nombre_figura_esperado(script: Path) -> str | None:
    """Intenta inferir el nombre de la figura de salida a partir del nombre
    del script.  Ej.: figura_a1.py -> Figura_A1.png"""
    stem = script.stem  # e.g. "figura_a1"
    # Extraer letra y número
    m = re.match(r"figura_([a-h])(\d+(?:\.\d+)?)", stem, re.IGNORECASE)
    if m:
        letra = m.group(1).upper()
        numero = m.group(2)
        return f"Figura_{letra}{numero}.png"
    return None


# ---------------------------------------------------------------------------
# Ejecutar un script individual
# ---------------------------------------------------------------------------
def ejecutar_script(script: Path) -> dict:
    """Ejecuta un script y devuelve un dict con el resultado."""
    resultado = {
        "script": script,
        "nombre": script.name,
        "carpeta": script.parent.name,
        "placeholder": False,
        "exito": False,
        "genero_grafica": False,
        "error": "",
        "stdout": "",
        "duracion": 0.0,
    }

    # Saltar placeholders
    if es_placeholder(script):
        resultado["placeholder"] = True
        resultado["exito"] = True  # No falló, simplemente no hace nada
        return resultado

    # Snapshot  de output antes
    antes = archivos_output_antes()

    inicio = time.time()
    try:
        proc = subprocess.run(
            [sys.executable, str(script)],
            capture_output=True,
            text=True,
            timeout=300,  # 5 min máx por script
            cwd=str(script.parent),
            env={**os.environ, "ANUARIO_DISABLE_AUTO_PRINT": "1"},
        )
        resultado["duracion"] = time.time() - inicio
        resultado["stdout"] = proc.stdout
        resultado["error"] = proc.stderr

        if proc.returncode == 0:
            resultado["exito"] = True
        else:
            resultado["exito"] = False
            if not resultado["error"]:
                resultado["error"] = f"Código de salida: {proc.returncode}"

    except subprocess.TimeoutExpired:
        resultado["duracion"] = time.time() - inicio
        resultado["exito"] = False
        resultado["error"] = "TIMEOUT: el script excedió 5 minutos."
    except Exception as e:
        resultado["duracion"] = time.time() - inicio
        resultado["exito"] = False
        resultado["error"] = str(e)

    # Verificar si generó una gráfica nueva
    despues = archivos_output_antes()
    nuevos = despues - antes
    if nuevos:
        resultado["genero_grafica"] = True
    else:
        # Verificar si el archivo esperado ya existía y fue actualizado
        esperado = nombre_figura_esperado(script)
        if esperado and (OUTPUT_DIR / esperado).exists():
            # Chequear si el timestamp de modificación es reciente
            mtime = (OUTPUT_DIR / esperado).stat().st_mtime
            if mtime >= inicio:
                resultado["genero_grafica"] = True

    return resultado


# ---------------------------------------------------------------------------
# Generar reporte Word
# ---------------------------------------------------------------------------
def generar_reporte_word(resultados: list[dict], duracion_total: float) -> Path:
    """Genera un documento Word con el resumen de ejecución."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print(f"\n{Color.AMARILLO}⚠  Instalando python-docx...{Color.RESET}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ── Título ──
    titulo_p = doc.add_heading("Reporte de Ejecución — Anuario Estadístico 2024", level=0)
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadatos ──
    ahora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    doc.add_paragraph(f"Fecha de ejecución: {ahora}")
    doc.add_paragraph(f"Duración total: {duracion_total:.1f} segundos")
    doc.add_paragraph(f"Total de scripts: {len(resultados)}")

    # Clasificar
    exitosos      = [r for r in resultados if r["exito"] and not r["placeholder"]]
    con_error     = [r for r in resultados if not r["exito"]]
    placeholders  = [r for r in resultados if r["placeholder"]]
    sin_grafica   = [r for r in resultados if r["exito"] and not r["placeholder"] and not r["genero_grafica"]]
    con_grafica   = [r for r in resultados if r["exito"] and not r["placeholder"] and r["genero_grafica"]]

    # ── Resumen general ──
    doc.add_heading("Resumen General", level=1)
    tabla_resumen = doc.add_table(rows=6, cols=2)
    tabla_resumen.style = 'Light Shading Accent 1'
    tabla_resumen.alignment = WD_TABLE_ALIGNMENT.CENTER

    filas_resumen = [
        ("Total de scripts", str(len(resultados))),
        ("✅ Completados con gráfica", str(len(con_grafica))),
        ("⚠️ Completados sin gráfica", str(len(sin_grafica))),
        ("❌ Con error", str(len(con_error))),
        ("📄 Placeholders (vacíos)", str(len(placeholders))),
        ("⏱️ Duración total", f"{duracion_total:.1f} s"),
    ]
    for i, (clave, valor) in enumerate(filas_resumen):
        tabla_resumen.cell(i, 0).text = clave
        tabla_resumen.cell(i, 1).text = valor

    # ── Detalle por carpeta ──
    doc.add_heading("Detalle por Carpeta", level=1)

    for carpeta in CARPETAS:
        scripts_carpeta = [r for r in resultados if r["carpeta"] == carpeta]
        if not scripts_carpeta:
            continue

        doc.add_heading(f"Carpeta {carpeta}", level=2)

        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = 'Light Grid Accent 1'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        header = tabla.rows[0].cells
        header[0].text = "Script"
        header[1].text = "Estado"
        header[2].text = "Gráfica"
        header[3].text = "Duración"

        for h in header:
            for p in h.paragraphs:
                for run in p.runs:
                    run.bold = True

        for r in scripts_carpeta:
            row = tabla.add_row().cells
            row[0].text = r["nombre"]

            if r["placeholder"]:
                row[1].text = "Placeholder"
                row[2].text = "—"
            elif r["exito"]:
                row[1].text = "✅ OK"
                row[2].text = "Sí" if r["genero_grafica"] else "No"
            else:
                row[1].text = "❌ Error"
                row[2].text = "No"

            row[3].text = f"{r['duracion']:.1f}s" if r["duracion"] > 0 else "—"

    # ── Scripts con error (detalle) ──
    if con_error:
        doc.add_heading("Detalle de Errores", level=1)
        for r in con_error:
            doc.add_heading(f"{r['carpeta']}/{r['nombre']}", level=3)
            # Mostrar las últimas líneas del error
            error_text = r["error"].strip()
            if error_text:
                lineas = error_text.split('\n')
                # Mostrar últimas 20 líneas
                lineas_mostrar = lineas[-20:] if len(lineas) > 20 else lineas
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run('\n'.join(lineas_mostrar))
                run.font.size = Pt(8)
                run.font.name = 'Consolas'
                run.font.color.rgb = RGBColor(180, 0, 0)

    # ── Scripts sin gráfica ──
    if sin_grafica:
        doc.add_heading("Scripts que no generaron gráfica", level=1)
        for r in sin_grafica:
            doc.add_paragraph(f"• {r['carpeta']}/{r['nombre']}", style='List Bullet')

    # ── Guardar ──
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"reporte_ejecucion_{timestamp}.docx"
    ruta_reporte = REPORT_DIR / nombre_archivo
    doc.save(str(ruta_reporte))

    return ruta_reporte


# ---------------------------------------------------------------------------
# Resumen en terminal
# ---------------------------------------------------------------------------
def imprimir_resumen(resultados: list[dict], duracion_total: float, ruta_reporte: Path) -> None:
    """Imprime un resumen colorido en la terminal."""

    exitosos      = [r for r in resultados if r["exito"] and not r["placeholder"]]
    con_error     = [r for r in resultados if not r["exito"]]
    placeholders  = [r for r in resultados if r["placeholder"]]
    sin_grafica   = [r for r in resultados if r["exito"] and not r["placeholder"] and not r["genero_grafica"]]
    con_grafica   = [r for r in resultados if r["exito"] and not r["placeholder"] and r["genero_grafica"]]

    titulo("RESUMEN DE EJECUCIÓN")

    print(f"  Total de scripts:  {Color.BOLD}{len(resultados)}{Color.RESET}")
    print(f"  Duración total:    {Color.BOLD}{duracion_total:.1f}s{Color.RESET}")
    print()

    # ── Archivos completados con gráfica ──
    seccion(f"✅ COMPLETADOS CON GRÁFICA ({len(con_grafica)})")
    if con_grafica:
        for r in con_grafica:
            print(f"  {Color.VERDE}✔ {r['carpeta']}/{r['nombre']}{Color.RESET}"
                  f"  {Color.GRIS}({r['duracion']:.1f}s){Color.RESET}")
    else:
        print(f"  {Color.GRIS}(ninguno){Color.RESET}")

    # ── Archivos con error ──
    seccion(f"❌ CON ERROR ({len(con_error)})")
    if con_error:
        for r in con_error:
            print(f"  {Color.ROJO}✘ {r['carpeta']}/{r['nombre']}{Color.RESET}")
            # Mostrar última línea del error
            err_lines = r["error"].strip().split('\n')
            ultima = err_lines[-1] if err_lines else ""
            if ultima:
                print(f"    {Color.GRIS}└─ {ultima[:120]}{Color.RESET}")
    else:
        print(f"  {Color.GRIS}(ninguno){Color.RESET}")

    # ── Scripts que no generan gráfica ──
    seccion(f"⚠️  SIN GRÁFICA ({len(sin_grafica)})")
    if sin_grafica:
        for r in sin_grafica:
            print(f"  {Color.AMARILLO}◦ {r['carpeta']}/{r['nombre']}{Color.RESET}")
    else:
        print(f"  {Color.GRIS}(ninguno){Color.RESET}")

    # ── Placeholders ──
    seccion(f"📄 PLACEHOLDERS / VACÍOS ({len(placeholders)})")
    if placeholders:
        for r in placeholders:
            print(f"  {Color.GRIS}· {r['carpeta']}/{r['nombre']}{Color.RESET}")
    else:
        print(f"  {Color.GRIS}(ninguno){Color.RESET}")

    # ── Ruta del reporte ──
    print(f"\n{Color.CYAN}{Color.BOLD}📝 Reporte Word guardado en:{Color.RESET}")
    print(f"   {ruta_reporte}\n")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    titulo("ORQUESTADOR — Anuario Estadístico IFT 2024")
    print(f"  Carpetas:  {', '.join(CARPETAS)}")
    print(f"  Output:    {OUTPUT_DIR}")
    print(f"  Reportes:  {REPORT_DIR}")

    # Asegurar que la carpeta output existe
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    scripts = recopilar_scripts()
    total = len(scripts)
    print(f"\n  Scripts encontrados: {Color.BOLD}{total}{Color.RESET}\n")

    resultados = []
    inicio_total = time.time()

    for i, script in enumerate(scripts, 1):
        etiqueta = f"{script.parent.name}/{script.name}"
        es_ph = es_placeholder(script)

        if es_ph:
            status = f"{Color.GRIS}[PLACEHOLDER]{Color.RESET}"
        else:
            status = f"{Color.CYAN}ejecutando...{Color.RESET}"

        print(f"  [{i:3d}/{total}] {etiqueta:40s} {status}", end="", flush=True)

        resultado = ejecutar_script(script)
        resultados.append(resultado)

        # Limpiar la línea y mostrar resultado
        if es_ph:
            print(f"\r  [{i:3d}/{total}] {etiqueta:40s} {Color.GRIS}[PLACEHOLDER]{Color.RESET}")
        elif resultado["exito"]:
            grafica_txt = f"{Color.VERDE}✔ gráfica{Color.RESET}" if resultado["genero_grafica"] else f"{Color.AMARILLO}sin gráfica{Color.RESET}"
            print(f"\r  [{i:3d}/{total}] {etiqueta:40s} {Color.VERDE}OK{Color.RESET} ({resultado['duracion']:.1f}s) {grafica_txt}")
        else:
            print(f"\r  [{i:3d}/{total}] {etiqueta:40s} {Color.ROJO}ERROR{Color.RESET} ({resultado['duracion']:.1f}s)")

    duracion_total = time.time() - inicio_total

    # Generar reporte Word
    print(f"\n{Color.CYAN}Generando reporte Word...{Color.RESET}")
    ruta_reporte = generar_reporte_word(resultados, duracion_total)

    # Resumen final
    imprimir_resumen(resultados, duracion_total, ruta_reporte)


if __name__ == "__main__":
    main()
